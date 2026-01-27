import os
import logging
from typing import List, Dict, Any, Optional
from pinecone import Pinecone, ServerlessSpec
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings

logger = logging.getLogger(__name__)

class RAGService:
    """Service for managing document preprocessing and Pinecone vector storage using Gemini Embeddings"""

    def __init__(self):
        self.pc = None
        self.index = None
        
        # Initialize Google Gemini for free embeddings
        if settings.GEMINI_API_KEY:
            genai.configure(api_key=settings.GEMINI_API_KEY)
        
        # Initialize Pinecone
        if settings.PINECONE_API_KEY:
            try:
                self.pc = Pinecone(api_key=settings.PINECONE_API_KEY)
                
                # Check if index exists, else create it
                if settings.PINECONE_INDEX_NAME not in self.pc.list_indexes().names():
                    self.pc.create_index(
                        name=settings.PINECONE_INDEX_NAME,
                        dimension=768, # Gemini text-embedding-004 dimension
                        metric='cosine',
                        spec=ServerlessSpec(
                            cloud='aws',
                            region='us-east-1'
                        )
                    )
                self.index = self.pc.Index(settings.PINECONE_INDEX_NAME)
            except Exception as e:
                logger.error(f"Failed to initialize Pinecone: {e}")

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file formats"""
        text = ""
        try:
            if file_path.endswith('.pdf') or 'pdf' in file_type:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif file_path.endswith('.docx') or 'vnd.openxmlformats-officedocument.wordprocessingml.document' in file_type:
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif file_path.endswith('.doc') or 'msword' in file_type:
                # Basic support for .doc via docx (not always compatible)
                try:
                    doc = Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                except:
                    text = "Legacy .doc format extraction not fully supported. Please convert to .docx or .pdf"
            elif file_path.endswith('.txt') or 'text/plain' in file_type:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                logger.warning(f"Unsupported file type for extraction: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            text = f"Error extracting text: {str(e)}"
        
        return text

    async def get_embeddings(self, texts: List[str], task_type: str = "retrieval_document") -> List[List[float]]:
        """Generate Gemini embeddings (free tier) for a list of texts"""
        if not settings.GEMINI_API_KEY:
            logger.error("Gemini API key not configured for embeddings")
            return []
        
        try:
            # Gemini models/text-embedding-004 is free and powerful
            result = genai.embed_content(
                model="models/text-embedding-004",
                content=texts,
                task_type=task_type
            )
            # result['embedding'] is a list of embeddings if multiple texts were passed
            return result['embedding']
        except Exception as e:
            logger.error(f"Failed to generate Gemini embeddings: {e}")
            return []

    async def index_document(self, agent_id: str, source_id: str, text: str, metadata: Dict[str, Any]):
        """Chunk, embed, and index a document in Pinecone using Gemini"""
        if not self.index:
            logger.warning("Pinecone index not available for indexing")
            return

        # 1. Chunk text
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100
        )
        chunks = text_splitter.split_text(text)
        
        if not chunks:
            return

        # 2. Get embeddings (Gemini)
        # Note: genai.embed_content handles batching internally
        embeddings = await self.get_embeddings(chunks, task_type="retrieval_document")
        if not embeddings:
            return

        # 3. Upsert to Pinecone
        vectors = []
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            vectors.append({
                "id": f"{source_id}_{i}",
                "values": emb,
                "metadata": {
                    **metadata,
                    "agent_id": agent_id,
                    "source_id": source_id,
                    "text": chunk, # Store text for retrieval
                    "name": metadata.get("name", "Document"), # Ensure name is in metadata
                    "chunk_idx": i
                }
            })
        
        try:
            # Batch upsert to Pinecone
            for i in range(0, len(vectors), 100):
                self.index.upsert(vectors=vectors[i:i+100], namespace=agent_id)
            logger.info(f"Indexed {len(vectors)} chunks for {source_id} in namespace {agent_id} using Gemini")
        except Exception as e:
            logger.error(f"Failed to upsert to Pinecone: {e}")

    async def delete_document(self, agent_id: str, source_id: str):
        """Delete all chunks for a document from Pinecone"""
        if not self.index: return
        try:
            self.index.delete(filter={"source_id": {"$eq": source_id}}, namespace=agent_id)
        except Exception as e:
            logger.error(f"Failed to delete document from Pinecone: {e}")

    async def query(self, agent_id: str, query_text: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Query Pinecone for relevant document chunks using Gemini query embeddings"""
        if not self.index or not settings.GEMINI_API_KEY:
            return []
        
        try:
            # 1. Get query embedding from Gemini
            resp = genai.embed_content(
                model="models/text-embedding-004",
                content=query_text,
                task_type="retrieval_query"
            )
            query_emb = resp['embedding']
            
            # 2. Query Pinecone
            results = self.index.query(
                namespace=agent_id,
                vector=query_emb,
                top_k=top_k,
                include_metadata=True
            )
            
            # 3. Filter and format results
            # Only include results with a decent similarity score (threshold: 0.35)
            relevant_chunks = []
            logger.info(f"Pinecone query for query_text='{query_text[:50]}...' in namespace '{agent_id}' returned {len(results.matches)} potential matches")
            
            for match in results.matches:
                logger.info(f"Match ID: {match.id}, Score: {match.score}")
                if match.score >= 0.35 and 'text' in match.metadata:
                    relevant_chunks.append({
                        "text": match.metadata['text'],
                        "source": match.metadata.get('name', 'Unknown Source'),
                        "score": match.score
                    })
            
            logger.info(f"Retrieved {len(relevant_chunks)} chunks above threshold 0.35")
            return relevant_chunks
        except Exception as e:
            logger.error(f"Pinecone-Gemini query failed: {e}")
            return []

rag_service = RAGService()
